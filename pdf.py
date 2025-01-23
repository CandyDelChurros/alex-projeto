from docx import Document
import PyPDF2
import fitz
from PIL import Image
import io
import google.generativeai as genai

def save_images_from_pdf(pdf_file, output_folder):
    try:
        pdf_document = fitz.open(pdf_file)
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            images = page.get_images(full=True)

            for img_index, img in enumerate(images):
                xref = img[0]
                base_image = pdf_document.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]
                image = Image.open(io.BytesIO(image_bytes))
                output_path = f"{output_folder}/page_{page_num + 1}_img_{img_index + 1}.{image_ext}"
                image.save(output_path)
                print(f"Imagem salva em: {output_path}")

        print("Extração de imagens concluída!")
    except Exception as e:
        print(f"Erro ao extrair imagens: {e}")

def pdf_to_word_or_images(pdf_file, word_file, output_folder):
    try:
        has_images = False

        pdf_document = fitz.open(pdf_file)
        image_paths = []
        for page_num in range(len(pdf_document)):
            images = pdf_document[page_num].get_images(full=True)
            if images:
                has_images = True
                for img_index, img in enumerate(images):
                    xref = img[0]
                    base_image = pdf_document.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    output_path = f"{output_folder}/page_{page_num + 1}_img_{img_index + 1}.{image_ext}"
                    image_paths.append(output_path)

        if has_images:
            print("O PDF contém imagens. Iniciando a extração de imagens...")
            save_images_from_pdf(pdf_file, output_folder)

            genai.configure(api_key="KEY")
            model = genai.GenerativeModel("gemini-1.5-flash")

            doc = Document()
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'

            for image_path in image_paths:
                with open(image_path, "rb") as img_file:
                    img = Image.open(img_file)
                    response = model.generate_content([
                        "O que está escrito? Apenas transcreva o texto sem comentar nada.", img
                    ])
                    doc.add_paragraph(response.text)

            doc.save(word_file)
            print("Imagens processadas e conteúdo salvo no arquivo Word!")

        else:
            print("O PDF não contém imagens. Iniciando a conversão para Word...")
            with open(pdf_file, 'rb') as pdf:
                reader = PyPDF2.PdfReader(pdf)
                document = Document()

                for page in reader.pages:
                    text = page.extract_text()
                    document.add_paragraph(text)

                document.save(word_file)
                print(f"Conversão concluída! O arquivo Word foi salvo como: {word_file}")

    except Exception as e:
        print(f"Ocorreu um erro: {e}")

pdf_file = "insira seu pdf.pdf"  
word_file = "coloque o nome do doc.docx"  
output_folder = "imagens_extraidas"  

import os
os.makedirs(output_folder, exist_ok=True)

pdf_to_word_or_images(pdf_file, word_file, output_folder)
